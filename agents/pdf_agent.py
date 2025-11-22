"""
PDF-Rapport Agent

Skapar automatiska PDF-rapporter från GDPR-analys, audit-resultat, 
bristdetektion och sammanfattningar.
"""

from __future__ import annotations

import os
import tempfile
from typing import Dict, Any, List, Optional
from datetime import datetime
from pathlib import Path

from agents.gdpr_agent import GDPRReport
from agents.audit_agent import AuditReport
from rag.compliance_score import ComplianceScore
from rag.compliance_score import ComplianceScore


class PDFAgent:
    """Agent för att generera PDF-rapporter från olika analyser."""
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Args:
            output_dir: Katalog där PDF-filer ska sparas (default: temp directory)
        """
        self.output_dir = output_dir or tempfile.gettempdir()
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
    
    def generate_gdpr_report(self, report: GDPRReport, include_charts: bool = False) -> str:
        """
        Genererar PDF-rapport från GDPR-skanning.
        
        Args:
            report: GDPRReport från GDPR-agent
            include_charts: Om diagram ska inkluderas (kräver matplotlib)
            
        Returns:
            Path till genererad PDF-fil
        """
        # Skapa Word-dokument
        docx_path = self._create_gdpr_docx(report, include_charts)
        
        # Konvertera till PDF
        pdf_path = self._convert_docx_to_pdf(docx_path)
        
        # Ta bort temporär DOCX om PDF skapades
        if os.path.exists(pdf_path) and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except:
                pass
        
        return pdf_path
    
    def generate_audit_report(self, report: AuditReport, include_details: bool = True) -> str:
        """
        Genererar PDF-rapport från audit-skanning.
        
        Args:
            report: AuditReport från Audit-agent
            include_details: Om detaljerade findings ska inkluderas
            
        Returns:
            Path till genererad PDF-fil
        """
        # Skapa Word-dokument
        docx_path = self._create_audit_docx(report, include_details)
        
        # Konvertera till PDF
        pdf_path = self._convert_docx_to_pdf(docx_path)
        
        # Ta bort temporär DOCX om PDF skapades
        if os.path.exists(pdf_path) and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except:
                pass
        
        return pdf_path
    
    def generate_combined_report(
        self, 
        gdpr_report: Optional[GDPRReport] = None,
        audit_report: Optional[AuditReport] = None,
        title: str = "Kombinerad Analysrapport"
    ) -> str:
        """
        Genererar kombinerad PDF-rapport från flera analyser.
        
        Args:
            gdpr_report: GDPR-rapport (optional)
            audit_report: Audit-rapport (optional)
            title: Titel på rapporten
            
        Returns:
            Path till genererad PDF-fil
        """
        # Skapa Word-dokument
        docx_path = self._create_combined_docx(gdpr_report, audit_report, title)
        
        # Konvertera till PDF
        pdf_path = self._convert_docx_to_pdf(docx_path)
        
        # Ta bort temporär DOCX om PDF skapades
        if os.path.exists(pdf_path) and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except:
                pass
        
        return pdf_path
    
    def _create_gdpr_docx(self, report: GDPRReport, include_charts: bool = False) -> str:
        """Skapar Word-dokument från GDPR-rapport."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx krävs för PDF-generering. Installera med: pip install python-docx")
        
        doc = Document()
        
        # Titelblad
        title = doc.add_heading('GDPR-Compliance Rapport', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Dokument: {report.document_name}")
        doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # Sammanfattning
        doc.add_heading('Sammanfattning', level=1)
        doc.add_paragraph(report.summary)
        
        # Risk-score
        doc.add_heading('Riskbedömning', level=1)
        doc.add_paragraph(f"Riskpoäng: {report.risk_score}/100")
        doc.add_paragraph(f"Risknivå: {report.risk_level.upper()}")
        
        # Findings
        if report.findings:
            doc.add_heading('Identifierade Risker och Brister', level=1)
            
            for i, finding in enumerate(report.findings, 1):
                doc.add_heading(f'{i}. {finding.category}', level=2)
                doc.add_paragraph(f"Severitet: {finding.severity.upper()}")
                doc.add_paragraph(f"Plats: {finding.location}")
                doc.add_paragraph(f"Beskrivning: {finding.description}")
                doc.add_paragraph(f"Rekommendation: {finding.recommendation}")
                doc.add_paragraph("")  # Tom rad
        
        # Flaggor
        doc.add_heading('Identifierade Flaggor', level=1)
        flags = []
        if report.has_personnummer:
            flags.append("Personnummer hittades")
        if report.has_health_data:
            flags.append("Hälsodata hittades")
        if report.has_sensitive_categories:
            flags.append("Känsliga kategorier hittades")
        if report.missing_legal_basis:
            flags.append("Saknad rättslig grund")
        if report.missing_retention_period:
            flags.append("Saknad lagringsperiod")
        if report.missing_dpia:
            flags.append("Saknad DPIA")
        
        if flags:
            for flag in flags:
                doc.add_paragraph(f"• {flag}", style='List Bullet')
        else:
            doc.add_paragraph("Inga specifika flaggor identifierade.")
        
        # Spara
        filename = f"gdpr_report_{report.document_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _create_audit_docx(self, report: AuditReport, include_details: bool = True) -> str:
        """Skapar Word-dokument från Audit-rapport."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx krävs för PDF-generering. Installera med: pip install python-docx")
        
        doc = Document()
        
        # Titelblad
        title = doc.add_heading('Dokumentaudit Rapport', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Dokument: {report.document_name}")
        doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # Sammanfattning
        doc.add_heading('Sammanfattning', level=1)
        doc.add_paragraph(report.summary)
        
        # Statistik
        doc.add_heading('Statistik', level=1)
        doc.add_paragraph(f"Hög prioritet: {report.high_priority_count}")
        doc.add_paragraph(f"Medel prioritet: {report.medium_priority_count}")
        doc.add_paragraph(f"Låg prioritet: {report.low_priority_count}")
        doc.add_paragraph(f"Totalt: {len(report.findings)} findings")
        
        # Findings
        if report.findings and include_details:
            doc.add_heading('Detaljerade Findings', level=1)
            
            # Gruppera efter prioritet
            for priority in ["high", "medium", "low"]:
                priority_findings = [f for f in report.findings if f.priority == priority]
                if priority_findings:
                    doc.add_heading(f'{priority.upper()} Prioritet', level=2)
                    
                    for i, finding in enumerate(priority_findings, 1):
                        doc.add_heading(f'{i}. {finding.category}', level=3)
                        doc.add_paragraph(f"Plats: {finding.location}")
                        doc.add_paragraph(f"Problem: {finding.problem}")
                        doc.add_paragraph(f"Förklaring: {finding.explanation}")
                        doc.add_paragraph(f"Förslag: {finding.suggestion}")
                        doc.add_paragraph("")  # Tom rad
        
        # Spara
        filename = f"audit_report_{report.document_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _create_combined_docx(
        self, 
        gdpr_report: Optional[GDPRReport],
        audit_report: Optional[AuditReport],
        title: str
    ) -> str:
        """Skapar kombinerat Word-dokument från flera rapporter."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx krävs för PDF-generering. Installera med: pip install python-docx")
        
        doc = Document()
        
        # Titelblad
        title_heading = doc.add_heading(title, 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # GDPR-sektion
        if gdpr_report:
            doc.add_heading('GDPR-Analys', level=1)
            doc.add_paragraph(f"Dokument: {gdpr_report.document_name}")
            doc.add_paragraph(f"Riskpoäng: {gdpr_report.risk_score}/100 ({gdpr_report.risk_level})")
            doc.add_paragraph(gdpr_report.summary)
            doc.add_page_break()
        
        # Audit-sektion
        if audit_report:
            doc.add_heading('Dokumentaudit', level=1)
            doc.add_paragraph(f"Dokument: {audit_report.document_name}")
            doc.add_paragraph(audit_report.summary)
            doc.add_page_break()
        
        # Spara
        filename = f"combined_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def generate_compliance_report(
        self,
        document_name: str,
        compliance_score: ComplianceScore,
        gdpr_report: Optional[GDPRReport] = None,
        audit_report: Optional[AuditReport] = None,
        company_name: Optional[str] = None
    ) -> str:
        """
        Genererar superenkel V1 PDF-rapport från compliance-analys.
        
        V1: Minimal, funktional - ingen fancy layout ännu.
        
        Args:
            document_name: Namn på dokumentet som analyserats
            compliance_score: ComplianceScore från ComplianceScoreEngine
            gdpr_report: GDPR-rapport (optional)
            audit_report: Audit-rapport (optional)
            company_name: Företagsnamn för rapport (optional)
            
        Returns:
            Path till genererad PDF-fil
        """
        # Skapa Word-dokument
        docx_path = self._create_compliance_report_docx(
            document_name, compliance_score, gdpr_report, audit_report, company_name
        )
        
        # Konvertera till PDF
        pdf_path = self._convert_docx_to_pdf(docx_path)
        
        # Ta bort temporär DOCX om PDF skapades
        if os.path.exists(pdf_path) and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
            except:
                pass
        
        return pdf_path
    
    def _create_compliance_report_docx(
        self,
        document_name: str,
        compliance_score: ComplianceScore,
        gdpr_report: Optional[GDPRReport],
        audit_report: Optional[AuditReport],
        company_name: Optional[str]
    ) -> str:
        """Skapar superenkel Word-dokument från compliance-analys (V1)."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx krävs för PDF-generering. Installera med: pip install python-docx")
        
        doc = Document()
        
        # Titelblad (superenkel)
        title = doc.add_heading('Compliance-rapport', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if company_name:
            doc.add_paragraph(company_name)
        
        doc.add_paragraph(f"Dokument: {document_name}")
        doc.add_paragraph(f"Datum: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_page_break()
        
        # Scores (En sida)
        doc.add_heading('Compliance-scores', level=1)
        
        status_emoji = "🟢" if compliance_score.overall_compliance_score >= 80 else "🟡" if compliance_score.overall_compliance_score >= 60 else "🔴"
        status_text = "GRÖN" if compliance_score.overall_compliance_score >= 80 else "GUL" if compliance_score.overall_compliance_score >= 60 else "RÖD"
        
        doc.add_paragraph(f"Overall: {compliance_score.overall_compliance_score:.1f}/100 {status_emoji} {status_text}")
        doc.add_paragraph(f"GDPR Risk: {compliance_score.gdpr_risk_score:.1f}/100")
        doc.add_paragraph(f"Audit Quality: {compliance_score.audit_quality_score:.1f}/100")
        doc.add_paragraph(f"Completeness: {compliance_score.completeness_score:.1f}/100")
        
        doc.add_paragraph("")  # Tom rad
        doc.add_paragraph(compliance_score.summary)
        
        doc.add_page_break()
        
        # Top Findings (En sida)
        doc.add_heading('Top Findings och Rekommendationer', level=1)
        
        # GDPR Top 3
        if gdpr_report and gdpr_report.findings:
            doc.add_heading('GDPR-risker', level=2)
            high_findings = [f for f in gdpr_report.findings if f.severity == "high"]
            medium_findings = [f for f in gdpr_report.findings if f.severity == "medium"]
            top_gdpr = (high_findings + medium_findings)[:3]
            
            for i, finding in enumerate(top_gdpr, 1):
                doc.add_paragraph(f"{i}. {finding.category.upper()} ({finding.severity})")
                doc.add_paragraph(f"   {finding.description}")
                doc.add_paragraph(f"   💡 {finding.recommendation}")
                doc.add_paragraph("")
        
        # Audit Top 3-5
        if audit_report and audit_report.findings:
            doc.add_heading('Audit-findings', level=2)
            high_findings = [f for f in audit_report.findings if f.priority == "high"]
            medium_findings = [f for f in audit_report.findings if f.priority == "medium"]
            top_audit = (high_findings + medium_findings)[:5]
            
            for i, finding in enumerate(top_audit, 1):
                doc.add_paragraph(f"{i}. {finding.category.upper()} ({finding.priority} prioritet)")
                doc.add_paragraph(f"   Problem: {finding.problem}")
                doc.add_paragraph(f"   💡 {finding.suggestion}")
                doc.add_paragraph("")
        
        # Spara
        safe_name = "".join(c for c in document_name if c.isalnum() or c in "._-")[:50]
        filename = f"compliance_report_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """Konverterar DOCX till PDF."""
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Metod 1: Försök med docx2pdf (enkelt)
        try:
            import docx2pdf
            docx2pdf.convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return pdf_path
        except ImportError:
            pass
        except Exception as e:
            print(f"[PDF Agent] docx2pdf misslyckades: {e}")
        
        # Metod 2: Försök med LibreOffice (via command line)
        try:
            import subprocess
            # LibreOffice command: soffice --headless --convert-to pdf --outdir <dir> <file>
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', self.output_dir,
                docx_path
            ]
            result = subprocess.run(cmd, capture_output=True, timeout=30)
            if result.returncode == 0 and os.path.exists(pdf_path):
                return pdf_path
        except Exception as e:
            print(f"[PDF Agent] LibreOffice-konvertering misslyckades: {e}")
        
        # Fallback: returnera DOCX-fil om PDF-konvertering misslyckades
        print(f"[PDF Agent] Varning: Kunde inte konvertera till PDF. Returnerar DOCX: {docx_path}")
        print(f"[PDF Agent] Installera docx2pdf eller LibreOffice för PDF-stöd.")
        return docx_path

